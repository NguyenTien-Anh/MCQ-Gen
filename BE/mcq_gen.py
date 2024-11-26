
from llama_index.core.agent import ReActAgent
from llama_index.llms.openai import OpenAI
from llama_index.core.tools import QueryEngineTool, ToolMetadata
from llama_index.core import VectorStoreIndex, Document
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import Settings
from llama_index.core import PromptTemplate
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from docx import Document as DocxDocument
import json
import string


load_dotenv()

react_system_header_str = """\
Bạn được thiết kế để phục vụ một chức năng duy nhất: tạo ra các câu hỏi trắc nghiệm.\
Đáp án cuối cùng bạn cần đưa ra là một câu hỏi trắc nghiệm và chỉ định câu trả lời đúng.\

## Tools
Bạn có quyền truy cập vào nhiều công cụ khác nhau.
Bạn có trách nhiệm sử dụng các công cụ theo trình tự bất kỳ mà bạn cho là phù hợp để hoàn thành nhiệm vụ trước mắt.
Điều này có thể yêu cầu chia nhiệm vụ thành các nhiệm vụ phụ và sử dụng các công cụ khác nhau để hoàn thành từng nhiệm vụ phụ.
Bạn có quyền truy cập vào các công cụ sau:
{tool_desc}

## Output Format
Để tạo hoặc đánh giá các câu hỏi trắc nghiệm, vui lòng sử dụng định dạng sau.

```
Thought: Tôi cần sử dụng một công cụ để giúp tôi tạo hoặc đánh giá các câu hỏi trắc nghiệm đảm bảo đầu ra theo yêu cầu.
Action: tên công cụ (một trong {tool_names}) nếu sử dụng một công cụ.
Action Input: đầu vào của công cụ, ở định dạng JSON biểu thị kwargs (ví dụ: {{"input": "hello world", "num_beams": 5}})
```

Hãy LUÔN bắt đầu bằng một Suy nghĩ.

Vui lòng sử dụng định dạng JSON hợp lệ cho Đầu vào hành động. KHÔNG làm điều này {{'input': 'hello world', 'num_beams': 5}}.

Nếu định dạng này được sử dụng, người dùng sẽ phản hồi theo định dạng sau:

```
Observation: công cụ phản hồi

```

Bạn nên tiếp tục lặp lại định dạng trên cho đến khi có đủ thông tin để đáp ứng yêu cầu
mà không cần sử dụng thêm bất kỳ công cụ nào. Lúc đó bạn PHẢI trả lời ở một trong hai định dạng sau:

```
Thought: Tôi có thể đưa ra câu hỏi trắc nghiệm, mà không cần sử dụng thêm bất kỳ công cụ nào.
Answer: Câu trả lời phải là một câu hỏi trắc nghiệm với bốn lựa chọn và chỉ định câu trả lời đúng.  [câu trả lời ở đây]
```

```
Thought: Tôi không thể trả lời yêu cầu bằng các công cụ được cung cấp.
Answer: Xin lỗi, tôi không thể trả lời yêu cầu của bạn.
```

## Additional Rules
- ** Câu trả lời PHẢI là một câu hỏi trắc nghiệm và chỉ định câu trả lời đúng, được viết bằng tiếng việt **.
- ** Đảm bảo định dạng câu hỏi đầu ra là:
Câu hỏi: (Đưa ra câu hỏi)
A. Đáp án 1
B. Đáp án 2
C. Đáp án 3 (nếu có)
D. Đáp án 4 (nếu có)
E. Đáp án 5 (nếu có)
Đáp án đúng: (Chỉ ra đáp án) đáp án đúng 1, đáp án đúng 2 (nếu có),...**

---
HÃY CHẮC CHẮN CÂU HỎI TRẮC NGHIỆM ĐƯỢC VIẾT BẰNG TIẾNG VIỆT VÀ ĐÚNG ĐỊNH DẠNG TRÊN. 
## Current Conversation
Dưới đây là cuộc trò chuyện hiện tại bao gồm các tin nhắn của con người và trợ lý xen kẽ nhau.

"""
data = None

def get_bloom_evaluation(question):
    PROMPT_TEMPLATE_BLOOM = (
    "Dưới đây là một câu hỏi về một môn học bất kỳ."
    "\n -----------------------------\n"
    "{context_str}"
    "\n -----------------------------\n"
    "Bạn là một chuyên gia về giảng dạy, hãy đánh giá câu hỏi tôi đưa vào theo thang đo Bloom. Thang đo Bloom có 6 cấp độ như sau:"
    "\n\n1. **Nhớ (Remember)**: Cấp độ này yêu cầu người học ghi nhớ hoặc nhận diện thông tin đã học, không cần hiểu sâu hay giải thích thêm. "
    "Câu hỏi ở cấp độ này sẽ kiểm tra khả năng nhớ lại các khái niệm, sự kiện, hoặc định nghĩa. "
    "Ví dụ: 'Nêu định nghĩa của lực ma sát.'"
    "\n   A. Lực cản trở chuyển động của vật trên bề mặt\n   B. Lực giúp tăng tốc vật chuyển động\n   C. Lực hút giữa hai vật\n   D. Lực làm vật dừng lại hoàn toàn"
    
    "\n\n2. **Hiểu (Understand)**: Câu hỏi yêu cầu người học giải thích hoặc diễn giải ý nghĩa của thông tin đã học. "
    "Đây là cấp độ yêu cầu người học không chỉ nhớ thông tin mà còn phải hiểu và có thể giải thích được ý nghĩa của thông tin đó. "
    "Ví dụ: 'Tại sao một vật chuyển động lại dừng lại khi không có lực tác dụng thêm?'"
    "\n   A. Do lực hấp dẫn\n   B. Do lực ma sát làm tiêu hao động năng\n   C. Do trọng lượng vật giảm dần\n   D. Do vận tốc ban đầu không đủ lớn"
    
    "\n\n3. **Áp dụng (Apply)**: Câu hỏi yêu cầu sử dụng kiến thức trong tình huống thực tế. "
    "Ở cấp độ này, người học phải vận dụng các lý thuyết hoặc nguyên lý vào một tình huống mới hoặc thực tế. "
    "Ví dụ: 'Một xe ô tô đang di chuyển với vận tốc 60 km/h, lực ma sát tác dụng lên xe là 500 N. Tính lực kéo cần thiết để xe giữ nguyên tốc độ.'"
    "\n   A. 400 N\n   B. 500 N\n   C. 600 N\n   D. 0 N"
    
    "\n\n4. **Phân tích (Analyze)**: Câu hỏi yêu cầu người học phân tích thông tin, chia nhỏ và xác định mối quan hệ giữa các phần của thông tin. "
    "Ở cấp độ này, người học phải phân tích các yếu tố hoặc mối quan hệ giữa các phần của vấn đề để hiểu rõ hơn về chúng. "
    "Ví dụ: 'Điều gì xảy ra khi tăng độ nhám của bề mặt tiếp xúc trong một hệ thống ma sát?'"
    "\n   A. Tăng ma sát và giảm chuyển động\n   B. Giảm ma sát và tăng chuyển động\n   C. Không thay đổi ma sát\n   D. Không ảnh hưởng đến chuyển động"
    
    "\n\n5. **Đánh giá (Evaluate)**: Câu hỏi yêu cầu người học đưa ra nhận định hoặc đánh giá về một vấn đề hoặc quan điểm. "
    "Câu hỏi ở cấp độ này yêu cầu người học đưa ra phán đoán dựa trên các tiêu chí hoặc bằng chứng, đánh giá một quan điểm hoặc giải pháp. "
    "Ví dụ: 'Đánh giá hiệu quả của việc sử dụng phanh ABS trên ô tô so với phanh thường.'"
    "\n   A. Giúp xe dừng nhanh hơn\n   B. Tăng độ an toàn khi phanh gấp\n   C. Giảm chi phí bảo trì xe\n   D. Không có lợi ích cụ thể nào"
    
    "\n\n6. **Sáng tạo (Create)**: Câu hỏi yêu cầu người học tạo ra một sản phẩm mới hoặc giải pháp sáng tạo từ kiến thức đã học. "
    "Đây là cấp độ yêu cầu người học phát triển một ý tưởng, thiết kế hoặc giải pháp mới. "
    "Ví dụ: 'Thiết kế một hệ thống phanh xe mới có khả năng tự động điều chỉnh lực phanh dựa trên điều kiện mặt đường.'"
    "\n   A. Hệ thống điều chỉnh lực bằng cảm biến nhiệt\n   B. Hệ thống ABS kết hợp với phân bố lực phanh điện tử\n   C. Phanh tay kết hợp với hệ thống phanh cơ khí\n   D. Phanh từ tính sử dụng lực hút nam châm"
    
    "\n\nHãy đưa ra câu trả lời là **1 level đánh giá** phù hợp nhất (Nhớ, Hiểu, Áp dụng, Phân tích, Đánh giá, hoặc Sáng tạo) mà không cần giải thích: {query_str}"
    )

    QA_PROMPT_BLOOM = PromptTemplate(PROMPT_TEMPLATE_BLOOM)
    query_engine_bloom = data.as_query_engine(similarity_top_k=3, text_qa_template=QA_PROMPT_BLOOM,
                                             llm=OpenAI(model='gpt-3.5-turbo-0125', temperature=0.1, max_tokens=512),
                                             max_tokens=-1)
    response = query_engine_bloom.query(question)
    return response



def select_topic(topic, quantity):
    PROMPT_TEMPLATE0_other = (
        "Dưới đây là một chủ đề về môn học."
        "\n-----------------------------\n"
        "{context_str}"
        "\n-----------------------------\n"
        "Bạn là một chuyên gia tạo câu hỏi trắc nghiệm, từ một chủ đề mà bạn tự chọn ra, hãy tìm ra các nội dung có thể sử dụng để tạo câu hỏi trắc nghiệm. "
        "Nếu không thể chọn đủ số lượng yêu cầu, hãy cố gắng chọn ra nhiều nhất có thể. "
        "Các nội dung được chọn phải có liên quan đến chủ đề đã chọn, được viết khái quát và ngắn gọn. "
        "Đảm bảo định dạng câu trả lời bao gồm các nội dung được viết theo kiểu list trong python. "
        "Ví dụ: ['nội dung 1', 'nội dung 2', 'nội dung 3',...]. "
        "Đưa ra danh sách nội dung: {query_str}"
    )

    QA_PROMPT0_other = PromptTemplate(PROMPT_TEMPLATE0_other)
    select_topic_prompt = ''
    query_engine0 = ''
    if topic != '':
        select_topic_prompt = "Hãy chọn " + str(
            quantity) + " nội dung liên quan đến chủ đề \"" + topic + "\" và đưa ra duy nhất một câu trả lời đúng định dạng kiểu list trong python."
        query_engine0 = data.as_query_engine(similarity_top_k=3, text_qa_template=QA_PROMPT0_other,
                                             llm=OpenAI(model='gpt-3.5-turbo-0125', temperature=0.1, max_tokens=512),
                                             max_tokens=-1)
    else:
        select_topic_prompt = "Hãy chọn " + str(
            quantity) + " nội dung bất kì trong dữ liệu bạn có và đưa ra duy nhất một câu trả lời đúng định dạng kiểu list trong python."
        query_engine0 = data.as_query_engine(similarity_top_k=3, text_qa_template=QA_PROMPT0_other,
                                             llm=OpenAI(model='gpt-3.5-turbo-0125', temperature=0.1, max_tokens=512),
                                             max_tokens=-1)

    response = query_engine0.query(select_topic_prompt)
    subTopics = str(response).split('[')[-1][:-1]
    return subTopics


def check(s):
    if s.startswith("Câu") and s.find("A.") != -1 and s.find("B.") != -1 and s.lower().find("đáp án") != -1:
        return True
    return False


def read_pdf_file(file):
    file_content = ''
    reader = PdfReader(file)
    num_pages = len(reader.pages)
    for page_num in range(num_pages):
        page_object = reader.pages[page_num]
        text = page_object.extract_text()
        file_content += " " + text
    return file_content


def read_docx_file(file):
    file_content = ''
    doc = DocxDocument(file)
    for para in doc.paragraphs:
        file_content += " " + para.text
    return file_content


def read_txt_file(file):
    file_content = file.read().decode('utf-8')
    return file_content


def mcqGen(topic, quantity, difficulty, file, inputText, status, type, number_of_answers=4, is_check=True):
    print("TOPIC: ", topic)
    print("QUANTITY: ", quantity)
    print("DIFFICULTY: ", difficulty)
    print("STATUS: ", status)
    print("TYPE: ", type)
    if is_check:
        print("CÓ KIỂM TRA LẠI !!!")
        return mcqGen_with_check(topic, quantity, difficulty, file, inputText, status, type, number_of_answers)
    else:
        print("KHÔNG KIỂM TRA LẠI !!!")
        return mcqGen_without_check(topic, quantity, difficulty, file, inputText, status, type, number_of_answers)


def format_mcq(mcqs):
    def parse_question(mcq):
        print("**CÂU HỎI TRƯỚC KHI FORMAT****************************")
        print(mcq)
        print("******************************************************")
        lines = mcq.split("\n")
        question = lines[0].strip()
        answers = []
        try:
            correct_answer = lines[-1].split(": ")[1].strip()
        except:
            correct_answer = ""

        print("ĐÁP ÁN ĐÚNG: ", correct_answer)

        for line in lines[1:-1]:
            answer_text = line.strip(string.punctuation)

            if answer_text == "":
                continue

            is_correct = "true" if (correct_answer.lower().find(answer_text.lower()) != -1 or
                                    answer_text.lower().find(correct_answer.lower()) != -1) else "false"
            print(f"CÂU TRẢ LỜI: {answer_text} - IS_CORRECT: {is_correct}")
            answers.append({"answer": answer_text, "isCorrectAnswer": is_correct})

        question_dict = {
            "question": question,
            "answers": answers
        }
        return question_dict

    return [parse_question(mcq) for mcq in mcqs]



def mcqGen_with_check(topic, quantity, difficulty, file, inputText, status, type, number_of_answers=4):
    print("NUM ANSWERS: ", number_of_answers)
    global data, file_content

    if status == 'true':
        print('ĐANG TẠO DATA ...')
        # file_content = ""
        # if file is not None:
            # print("ĐANG ĐỌC FILE ...")
            # ext_file = file.filename.split('.')[-1]
            # if ext_file == 'pdf':
                # file_content = read_pdf_file(file)
            # elif ext_file == 'docx':
                # file_content = read_docx_file(file)
            # elif ext_file == 'txt':
                # file_content = read_txt_file(file)
            # else:
                # raise ValueError("Unsupported file type")
            # print("ĐỌC FILE THÀNH CÔNG !!!")

        text_splitter = SentenceSplitter(chunk_size=512, chunk_overlap=10)
        content = file_content if file is not None else inputText
        gpt_documents = [Document(text=content)]
        data = VectorStoreIndex.from_documents(documents=gpt_documents, transformations=[text_splitter])
        print("TẠO DATA THÀNH CÔNG !!!")

    print('ĐANG TẠO CÂU HỎI ...')
    llm = OpenAI(model="gpt-3.5-turbo-0125")
    text_splitter = SentenceSplitter(chunk_size=512, chunk_overlap=10)
    Settings.text_splitter = text_splitter
    bloom_dict = {
    "Nhớ": "Câu hỏi yêu cầu người học ghi nhớ hoặc nhận diện thông tin đã học trước đó. Câu hỏi chỉ yêu cầu người học có thể nhớ lại các sự kiện, khái niệm, thuật ngữ, hoặc định nghĩa mà họ đã học. Câu hỏi ở cấp độ này chỉ yêu cầu nhớ lại thông tin, không yêu cầu giải thích hay phân tích gì thêm. Ví dụ: 'Đâu là năm diễn ra Cách mạng Tháng Tám ở Việt Nam?' \nA. 1945 \nB. 1954 \nC. 1975 \nD. 1986",
    
    "Hiểu": "Câu hỏi yêu cầu người học giải thích hoặc diễn giải ý nghĩa của thông tin đã học. Người học phải hiểu và nắm vững ý nghĩa của thông tin trước khi có thể diễn đạt lại bằng từ ngữ của mình. Câu hỏi này yêu cầu người học phải làm rõ những gì họ đã học thay vì chỉ đơn giản là nhớ thông tin. Ví dụ: 'Chọn câu trả lời đúng nhất để giải thích tại sao lá cây có màu xanh?' \nA. Do chứa diệp lục hấp thụ ánh sáng xanh \nB. Do chứa diệp lục phản xạ ánh sáng xanh \nC. Do chứa nước trong tế bào lá \nD. Do chứa các sắc tố hấp thụ tất cả ánh sáng ngoại trừ xanh",
    
    "Áp dụng": "Câu hỏi yêu cầu sử dụng kiến thức đã học trong các tình huống thực tế. Người học cần phải áp dụng các lý thuyết hoặc nguyên lý vào một tình huống mới. Đây là cấp độ yêu cầu người học sử dụng các công cụ hoặc quy tắc đã học để giải quyết vấn đề. Ví dụ: 'Nếu một tam giác có hai cạnh là 3 cm và 4 cm, đâu là độ dài cạnh huyền?' \nA. 5 cm \nB. 6 cm \nC. 7 cm \nD. 8 cm",
    
    "Phân tích": "Câu hỏi yêu cầu người học phân tích thông tin, chia nhỏ thành các phần và hiểu mối quan hệ giữa chúng. Người học cần phải xem xét các yếu tố chi tiết và hiểu cách chúng liên kết hoặc tác động với nhau. Đây là cấp độ đòi hỏi tư duy phức tạp và khả năng phân tích sâu sắc. Ví dụ: 'Trong bài thơ “Tây Tiến” của Quang Dũng, chi tiết nào sau đây thể hiện tinh thần hào hùng của người lính?' \nA. 'Sông Mã xa rồi Tây Tiến ơi!' \nB. 'Đêm mơ Hà Nội dáng kiều thơm' \nC. 'Rải rác biên cương mồ viễn xứ' \nD. 'Áo bào thay chiếu anh về đất'",
    
    "Đánh giá": "Câu hỏi yêu cầu đưa ra phán đoán hoặc nhận xét về một ý tưởng, quan điểm dựa trên tiêu chí hoặc bằng chứng học được. Người học cần phải sử dụng lý thuyết và các dữ liệu có sẵn để đánh giá một vấn đề hoặc giải pháp. Đây là cấp độ yêu cầu đưa ra quan điểm cá nhân dựa trên các lý lẽ vững chắc. Ví dụ: 'Đánh giá ý kiến sau: Biến đổi khí hậu là thách thức lớn nhất đối với nhân loại hiện nay. Bạn đồng ý với nhận định này không?' \nA. Hoàn toàn đồng ý \nB. Phần nào đồng ý \nC. Không đồng ý \nD. Không có ý kiến",
    
    "Sáng tạo": "Câu hỏi yêu cầu người học tạo ra một sản phẩm mới, ý tưởng mới hoặc giải pháp sáng tạo dựa trên những kiến thức đã học. Đây là cấp độ yêu cầu người học không chỉ tái tạo lại thông tin mà còn phải sáng tạo, phát triển những điều mới mẻ từ kiến thức hiện có. Ví dụ: 'Bạn được giao nhiệm vụ tổ chức một sự kiện tuyên truyền bảo vệ môi trường tại trường học. Đâu là ý tưởng phù hợp nhất?' \nA. Trồng cây xanh tại khuôn viên trường \nB. Tổ chức buổi hội thảo về môi trường \nC. Thiết kế áp phích tuyên truyền bảo vệ môi trường \nD. Cả A, B, và C đều đúng"
    }

    PROMPT_TEMPLATE_GEN = (
        "Bạn là một chuyên gia câu hỏi trắc nghiệm, hãy sinh ra câu hỏi trắc nghiệm trên nội dung đưa vào và chỉ ra đáp án đúng. "
        "Đầu vào là một nội dung môn học. "
        "Dữ liệu đưa vào là tài liệu về môn học."
        "\n-----------------------------\n"
        "{context_str}"
        "\n-----------------------------\n"
        "Quy trình thực hiện step by step để tạo câu hỏi trắc nghiệm."
        "\n{prompt_step_by_step}\n"
        "Ví dụ cho việc thực hiện quy trình là:"
        "\n{prompt_example}"
        "\n-----------------------------\n"
        "Bạn hãy thực hiện step by step quy trình trên và tạo 1 câu hỏi theo yêu cầu, đảm bảo định dạng câu trả lời:  {query_str}\n"
        "**Yêu cầu về định dạng câu trả lời là:**\n"
        "{attention}"
    )
    print(f"---------------PROMPT_TEMPLATE_GEN-----------------\n{PROMPT_TEMPLATE_GEN}")
    QA_PROMPT_GEN = PromptTemplate(PROMPT_TEMPLATE_GEN)
    with open('prompt.json', 'r', encoding='utf-8') as file:
        list_prompt_gen = json.load(file)

    gen_prompt_step_by_step = ""
    gen_prompt_example = ""
    gen_attention = ""
    for prompt_gen in list_prompt_gen['prompt_gen']:
        if prompt_gen['type'] == type and prompt_gen['number_of_answers'] == number_of_answers:
            gen_prompt_step_by_step = prompt_gen['prompt_step_by_step']
            gen_prompt_example = prompt_gen["prompt_example"]
            gen_attention = prompt_gen["attention"]

    print(f"\n\n\n------------gen_prompt_step_by_step--------------------\n{gen_prompt_step_by_step}")
    print(f"\n\n\n------------gen_prompt_example-------------------------\n{gen_prompt_example}")
    print(f"\n\n\n------------gen_attention------------------------------\n{gen_attention}")

    QA_PROMPT_GEN_FORMAT = QA_PROMPT_GEN.partial_format(prompt_step_by_step=gen_prompt_step_by_step,
                                                        prompt_example=gen_prompt_example, attention=gen_attention)
    # print(QA_PROMPT_GEN_FORMAT)
    PROMPT_TEMPLATE_EVA = (
        "Bạn là một chuyên gia về câu hỏi trắc nghiệm, hãy kiểm tra lại độ chính xác của câu hỏi và chỉnh sửa lại chúng tốt hơn. "
        "Nếu độ khó không đạt yêu cầu thì thực hiện chỉnh sửa lại độ khó. "
        "Mức độ khó yêu cầu là:" 
        "{difficulty_bloom}"
        "Đầu vào là 1 câu hỏi trắc nghiệm về môn học. "
        "Dữ liệu đưa vào là tài liệu về môn học."
        "\n-----------------------------\n"
        "{context_str}"
        "\n-----------------------------\n"
        "Hãy thực hiện step by step các bước theo quy trình để đánh giá câu hỏi trắc nghiệm."
        "\n{prompt_step_by_step}\n"
        "Ví dụ cho cách thực hiện quy trình trên.\n"
        "{prompt_example}"
        "\n-----------------------------\n"
        "Định dạng của câu hỏi cần đánh giá là:\n"
        "**{attention_eva}**"
        "\n-----------------------------\n"
        "Đầu ra là 1 lời đánh giá và một câu hỏi trắc nghiệm. Hãy đánh giá và cập nhật câu hỏi:  {query_str}. "
        "Đảm bảo định dạng phản hồi là 1 lời đánh giá và 1 câu hỏi trắc nghiệm, sau đó chỉ rõ đáp án đúng."
    )
    print(f"\n\n\n---------------PROMPT_TEMPLATE_EVA-----------------\n{PROMPT_TEMPLATE_EVA}")
    QA_PROMPT_EVA = PromptTemplate(PROMPT_TEMPLATE_EVA)
    eva_prompt_step_by_step = "Bước 1: Kiểm tra độ rõ ràng của câu hỏi và các đáp án.\nBước 2: Đánh giá tính liên quan của câu hỏi đến nội dung đã học.\nBước 3: Xác minh tính đúng đắn của đáp án đúng và sai.\nBước 4: Đánh giá độ khó của câu hỏi.\nBước 5: Đánh giá tính khách quan của câu hỏi.\nBước 6: Sửa đổi câu hỏi nếu cần thiết dựa trên các đánh giá.\nBước 7: Trả về câu hỏi đã được sửa đổi hoặc câu hỏi ban đầu."
    eva_prompt_example = "Bước 1: Câu hỏi \"Trong các loại khóa sau, loại nào được sử dụng để xác định duy nhất một bản ghi trong bảng?\" là rõ ràng và dễ hiểu.\nBước 2: Câu hỏi này liên quan trực tiếp đến nội dung đã học về cơ sở dữ liệu và các loại khóa.\nBước 3: Đáp án A) Khóa chính (Primary Key) là đúng. Các đáp án B) Khóa ngoại (Foreign Key), C) Khóa duy nhất (Unique Key), D) Khóa kết hợp (Composite Key), và E) Khóa tạm thời (Temporary Key) đều không phải là đáp án đúng cho câu hỏi này.\nBước 4: Độ khó của câu hỏi được đánh giá là trung bình vì nó yêu cầu người học hiểu rõ về các loại khóa trong cơ sở dữ liệu.\nBước 5: Câu hỏi này là khách quan, không thiên vị và không dẫn đến bất kỳ sự nhầm lẫn nào cho người trả lời.\nBước 6: Không cần sửa đổi câu hỏi, vì nó đã đạt yêu cầu và rõ ràng.\nCâu hỏi trắc nghiệm cuối cùng: \"Trong các loại khóa sau, loại nào được sử dụng để xác định duy nhất một bản ghi trong bảng?\" A) Khóa chính (Primary Key) B) Khóa ngoại (Foreign Key) C) Khóa duy nhất (Unique Key) D) Khóa kết hợp (Composite Key) E) Khóa tạm thời (Temporary Key) Đáp án đúng: A) Khóa chính (Primary Key)."
    attention_eva_dict = {
        "MultipleChoice": "Câu hỏi trắc nghiệm MultipleChoice gồm " + str(number_of_answers) + " đáp án và có ít nhất 2 đáp án đúng.",
        "SingleChoice": "Câu hỏi trắc nghiệm SingleChoice gồm " + str(number_of_answers) + " đáp án và có 1 đáp án đúng, " + str(number_of_answers - 1) + " đáp án sai.",
        "TrueFalse": "Câu hỏi trắc nghiệm TrueFalse chỉ gồm 2 loại đáp án là 'đúng' và 'sai'.",
    }
    print(f"\n\n\n---------------eva_prompt_step_by_step-------------------\n{eva_prompt_step_by_step}")
    print(f"\n\n\n---------------eva_prompt_example----------------------\n{eva_prompt_example}")
    print(f"\n\n\n---------------attention_eva_dict[type]----------------\n{attention_eva_dict[type]}")
    QA_PROMPT_EVA_FORMAT = QA_PROMPT_EVA.partial_format(difficulty_bloom=bloom_dict[difficulty],
                                                        prompt_step_by_step=eva_prompt_step_by_step,
                                                        prompt_example=eva_prompt_example,
                                                        attention_eva=attention_eva_dict[type])
    # print(QA_PROMPT_EVA_FORMAT)

    # GPT
    query_engine1 = data.as_query_engine(similarity_top_k=3, text_qa_template=QA_PROMPT_GEN_FORMAT,
                                         llm=OpenAI(model='gpt-3.5-turbo-0125', temperature=0.5, max_tokens=512),
                                         max_tokens=-1)
    query_engine2 = data.as_query_engine(similarity_top_k=3, text_qa_template=QA_PROMPT_EVA_FORMAT,
                                         llm=OpenAI(model='gpt-3.5-turbo-0125', temperature=0.1, max_tokens=512),
                                         max_tokens=-1)

    query_engine_tools = [
        QueryEngineTool(
            query_engine=query_engine1,
            metadata=ToolMetadata(
                name="Create",
                description=(
                    "Đầu vào là một yêu cầu. Đầu ra là một câu hỏi trắc nghiệm và có chỉ rõ đáp án đúng. "
                    "Tạo câu hỏi trắc nghiệm về nội dung được yêu cầu. "
                    "Sử dụng các công cụ khác để đánh giá câu hỏi."
                ),
            ),
        ),
        QueryEngineTool(
            query_engine=query_engine2,
            metadata=ToolMetadata(
                name="Check1",
                description=(
                    "Đầu vào là một câu hỏi trắc nghiệm. Đầu ra là 1 câu đánh giá và 1 câu hỏi trắc nghiệm. Hãy chỉ rõ câu trả lời đúng.  "
                    "Tiến hành đánh giá câu hỏi. Giải thích câu trả lời đúng, nếu câu hỏi hoặc câu trả lời sai thì thực hiện chỉnh sửa lại. "
                    "Nếu độ khó không đạt yêu cầu thì thực hiện chỉnh sửa lại độ khó. "
                    "Nếu không có câu trả lời đúng thì hãy sửa lại câu trả lời. "
                    "Nếu các đáp án tương tự nhau thì hãy sửa lại. "
                    "Cải thiện câu hỏi trắc nghiệm. "
                    "Kết quả cuối cùng là 1 câu hỏi trắc nghiệm."
                ),
            ),
        ),
        QueryEngineTool(
            query_engine=query_engine2,
            metadata=ToolMetadata(
                name="Check2",
                description=(
                    "Đầu vào là một câu hỏi trắc nghiệm. Đầu ra là 1 câu đánh giá và 1 câu hỏi trắc nghiệm. Hãy chỉ rõ câu trả lời đúng. "
                    "Tiến hành đánh giá câu hỏi. Giải thích câu trả lời đúng, nếu câu hỏi hoặc câu trả lời sai thì thực hiện chỉnh sửa lại. "
                    "Nếu độ khó không đạt yêu cầu thì thực hiện chỉnh sửa lại độ khó. "
                    "Nếu không có câu trả lời đúng thì hãy sửa lại câu trả lời. "
                    "Nếu các đáp án tương tự nhau thì hãy sửa lại. "
                    "Cải thiện câu hỏi trắc nghiệm. "
                    "Kết quả cuối cùng là 1 câu hỏi trắc nghiệm."
                ),
            ),
        ),
    ]

    agent = ReActAgent.from_tools(
        query_engine_tools,
        llm=llm,
        verbose=True,
    )

    react_system_prompt = PromptTemplate(react_system_header_str)
    agent.update_prompts({"agent_worker:system_prompt": react_system_prompt})
    agent.reset()
    subTopics = select_topic(topic, quantity)
    print(f"\n\n\n----------------subTopics----------------\n{subTopics}")
    list_topic = subTopics.split(", ")
    print(f"\n\n\n-----------------list_topic---------------\n{list_topic}")
    notify = ""
    if len(list_topic) < int(quantity):
        notify = "XIN LỖI CHÚNG TÔI KHÔNG THỂ SINH ĐỦ CÂU HỎI CHO CHỦ ĐỀ NÀY"
        print(notify)
    type_dict = {
        "MultipleChoice": "Tạo 1 câu hỏi trắc nghiệm MultipleChoice gồm " + str(
            number_of_answers) + " đáp án và có ít nhất 2 đáp án đúng.",
        "SingleChoice": "Tạo 1 câu hỏi trắc nghiệm singlechoice gồm " + str(
            number_of_answers) + " đáp án và có 1 đáp án đúng, " + str(number_of_answers - 1) + " đáp án sai.",
        "TrueFalse": "Tạo 1 câu hỏi trắc nghiệm TrueFalse chỉ gồm 2 loại đáp án đúng hoặc sai và có 1 đáp án đúng, 1 đáp án sai.",
    }
    mcqs = []
    for i in range(0, int(quantity)):
        if i > len(list_topic) - 1:
            continue
        genned_topic = list_topic[i]
        print(f"\n\n\n----------------------gened_topic-----------------\n{genned_topic}")
        if genned_topic[2] == '.':
            genned_topic = genned_topic[4:]
        if genned_topic[-2] == '.':
            genned_topic = genned_topic[:-2] + genned_topic[-1]
        kq = ""
        if topic != "":
            prompt = type_dict[
                     type] + " Câu hỏi có nội dung liên quan đến " + genned_topic + " trong chủ đề \"" + topic + "\". " + \
                 bloom_dict[difficulty]
        else:
            prompt = type_dict[
                         type] + " Câu hỏi có nội dung liên quan đến " + genned_topic + ". " + \
                     bloom_dict[difficulty]
        prompt = prompt + " Sau khi tạo câu hỏi sử dụng công cụ kiểm tra lại."
        print(f"\n\n\n-------------------prompt------------------------\n{prompt}")
        question = agent.chat(prompt)
        kq = str(question.response)
        eval_question = get_bloom_evaluation(kq)
        kq= str("topic: ")+str(topic)+"\n"+ str("difficulty fist: ")+str(difficulty)+"\n"+ str("eval_question_with_bloom: ")+str(eval_question)+"\n"+str("Câu hỏi trắc nghiệm: ")+str(kq)
        with open("kq_check.txt", "a", encoding="utf-8") as file:
            file.write(kq)
            file.write("\n\n\n")
        # mcqs.append(kq)
    # print("TẠO CÂU HỎI THÀNH CÔNG !!!")
    # print("ĐANG FORMAT CÂU HỎI ...")
    # mcqs = format_mcq(mcqs)
    # print("FORMAT CÂU HỎI THÀNH CÔNG !!!")
    # return mcqs, notify


def mcqGen_without_check(topic, quantity, difficulty, file, inputText, status, type, number_of_answers=4):
    print("NUM ANSWERS: ", number_of_answers)
    global data, file_content
    if file is None:
        print("File IS NONE. USE INPUT TEXT !!!")
    else:
        print("USING FILE !!!")
    if status == 'true':
        print('ĐANG TẠO DATA ...')
        # file_content = ""
        # if file is not None:
            # file_path = file
        # ext_file = file_path.split('.')[-1]
        # ext_file = file.filename.split('.')[-1]
        # if ext_file == 'pdf':
            # file_content = read_pdf_file(file)
            # elif ext_file == 'docx':
                # file_content = read_docx_file(file)
            # elif ext_file == 'txt':
                # file_content = read_txt_file(file)
            # else:
            # raise ValueError("Unsupported file type")
        # file_content = read_pdf_file(file)
        text_splitter = SentenceSplitter(chunk_size=512, chunk_overlap=10)
        content = file_content if file is not None else inputText
        gpt_documents = [Document(text=content)]
        data = VectorStoreIndex.from_documents(documents=gpt_documents, transformations=[text_splitter])
        print("TẠO DATA THÀNH CÔNG !!!")

    print('ĐANG TẠO CÂU HỎI ...')
    llm = OpenAI(model="gpt-3.5-turbo-0125")
    text_splitter = SentenceSplitter(chunk_size=512, chunk_overlap=10)
    Settings.text_splitter = text_splitter
    bloom_dict = {
    "Nhớ": "Câu hỏi yêu cầu người học ghi nhớ hoặc nhận diện thông tin đã học trước đó. Đây là cấp độ cơ bản nhất, chỉ yêu cầu người học có thể nhớ lại các sự kiện, khái niệm, thuật ngữ, hoặc định nghĩa mà họ đã học. Câu hỏi ở cấp độ này chỉ yêu cầu nhớ lại thông tin, không yêu cầu giải thích hay phân tích gì thêm. Ví dụ: 'Đâu là năm diễn ra Cách mạng Tháng Tám ở Việt Nam?' \nA. 1945 \nB. 1954 \nC. 1975 \nD. 1986",
    
    "Hiểu": "Câu hỏi yêu cầu người học giải thích hoặc diễn giải ý nghĩa của thông tin đã học. Người học phải hiểu và nắm vững ý nghĩa của thông tin trước khi có thể diễn đạt lại bằng từ ngữ của mình. Câu hỏi này yêu cầu người học phải làm rõ những gì họ đã học thay vì chỉ đơn giản là nhớ thông tin. Ví dụ: 'Chọn câu trả lời đúng nhất để giải thích tại sao lá cây có màu xanh?' \nA. Do chứa diệp lục hấp thụ ánh sáng xanh \nB. Do chứa diệp lục phản xạ ánh sáng xanh \nC. Do chứa nước trong tế bào lá \nD. Do chứa các sắc tố hấp thụ tất cả ánh sáng ngoại trừ xanh",
    
    "Áp dụng": "Câu hỏi yêu cầu sử dụng kiến thức đã học trong các tình huống thực tế. Người học cần phải áp dụng các lý thuyết hoặc nguyên lý vào một tình huống mới. Đây là cấp độ yêu cầu người học sử dụng các công cụ hoặc quy tắc đã học để giải quyết vấn đề. Ví dụ: 'Nếu một tam giác có hai cạnh là 3 cm và 4 cm, đâu là độ dài cạnh huyền?' \nA. 5 cm \nB. 6 cm \nC. 7 cm \nD. 8 cm",
    
    "Phân tích": "Câu hỏi yêu cầu người học phân tích thông tin, chia nhỏ thành các phần và hiểu mối quan hệ giữa chúng. Người học cần phải xem xét các yếu tố chi tiết và hiểu cách chúng liên kết hoặc tác động với nhau. Đây là cấp độ đòi hỏi tư duy phức tạp và khả năng phân tích sâu sắc. Ví dụ: 'Trong bài thơ “Tây Tiến” của Quang Dũng, chi tiết nào sau đây thể hiện tinh thần hào hùng của người lính?' \nA. 'Sông Mã xa rồi Tây Tiến ơi!' \nB. 'Đêm mơ Hà Nội dáng kiều thơm' \nC. 'Rải rác biên cương mồ viễn xứ' \nD. 'Áo bào thay chiếu anh về đất'",
    
    "Đánh giá": "Câu hỏi yêu cầu đưa ra phán đoán hoặc nhận xét về một ý tưởng, quan điểm dựa trên tiêu chí hoặc bằng chứng học được. Người học cần phải sử dụng lý thuyết và các dữ liệu có sẵn để đánh giá một vấn đề hoặc giải pháp. Đây là cấp độ yêu cầu đưa ra quan điểm cá nhân dựa trên các lý lẽ vững chắc. Ví dụ: 'Đánh giá ý kiến sau: Biến đổi khí hậu là thách thức lớn nhất đối với nhân loại hiện nay. Bạn đồng ý với nhận định này không?' \nA. Hoàn toàn đồng ý \nB. Phần nào đồng ý \nC. Không đồng ý \nD. Không có ý kiến",
    
    "Sáng tạo": "Câu hỏi yêu cầu người học tạo ra một sản phẩm mới, ý tưởng mới hoặc giải pháp sáng tạo dựa trên những kiến thức đã học. Đây là cấp độ yêu cầu người học không chỉ tái tạo lại thông tin mà còn phải sáng tạo, phát triển những điều mới mẻ từ kiến thức hiện có. Ví dụ: 'Bạn được giao nhiệm vụ tổ chức một sự kiện tuyên truyền bảo vệ môi trường tại trường học. Đâu là ý tưởng phù hợp nhất?' \nA. Trồng cây xanh tại khuôn viên trường \nB. Tổ chức buổi hội thảo về môi trường \nC. Thiết kế áp phích tuyên truyền bảo vệ môi trường \nD. Cả A, B, và C đều đúng"
    }

    PROMPT_TEMPLATE_GEN = (
        "Bạn là một chuyên gia câu hỏi trắc nghiệm, hãy sinh ra câu hỏi trắc nghiệm trên nội dung đưa vào và chỉ ra đáp án đúng. "
        "Đầu vào là một nội dung môn học. "
        "Dữ liệu đưa vào là tài liệu về môn học."
        "\n-----------------------------\n"
        "{context_str}"
        "\n-----------------------------\n"
        "Quy trình thực hiện step by step để tạo câu hỏi trắc nghiệm."
        "\n{prompt_step_by_step}\n"
        "Ví dụ cho việc thực hiện quy trình là:"
        "\n{prompt_example}"
        "\n-----------------------------\n"
        "Bạn hãy thực hiện step by step quy trình trên và tạo 1 câu hỏi theo yêu cầu, đảm bảo định dạng câu trả lời:  {query_str}\n"
        "**Yêu cầu về định dạng câu trả lời là:**\n"
        "{attention}"
    )
    print(f"---------------PROMPT_TEMPLATE_GEN-----------------\n{PROMPT_TEMPLATE_GEN}")
    QA_PROMPT_GEN = PromptTemplate(PROMPT_TEMPLATE_GEN)
    with open('prompt.json', 'r', encoding='utf-8') as file:
        list_prompt_gen = json.load(file)
    gen_prompt_step_by_step = ""
    gen_prompt_example = ""
    gen_attention = ""
    for prompt_gen in list_prompt_gen['prompt_gen']:
        if prompt_gen['type'] == type and prompt_gen['number_of_answers'] == number_of_answers:
            gen_prompt_step_by_step = prompt_gen['prompt_step_by_step']
            gen_prompt_example = prompt_gen["prompt_example"]
            gen_attention = prompt_gen["attention"]
    QA_PROMPT_GEN_FORMAT = QA_PROMPT_GEN.partial_format(prompt_step_by_step=gen_prompt_step_by_step,
                                                        prompt_example=gen_prompt_example, attention=gen_attention)

    print(f"\n\n\n------------gen_prompt_step_by_step--------------------\n{gen_prompt_step_by_step}")
    print(f"\n\n\n------------gen_prompt_example-------------------------\n{gen_prompt_example}")
    print(f"\n\n\n------------gen_attention------------------------------\n{gen_attention}")

    # print(QA_PROMPT_GEN_FORMAT)
    # GPT
    query_engine1 = data.as_query_engine(similarity_top_k=3, text_qa_template=QA_PROMPT_GEN_FORMAT,
                                         llm=OpenAI(model='gpt-3.5-turbo-0125', temperature=0.5, max_tokens=512),
                                         max_tokens=-1)
    subTopics = select_topic(topic, quantity)
    print(f"\n\n\n------------------subTopics----------------\n{subTopics}")
    list_topic = subTopics.split(", ")
    print(f"\n\n\n-------------------list_topic--------------\n{list_topic}")
    notify = ""
    if len(list_topic) < int(quantity):
        notify = "XIN LỖI CHÚNG TÔI KHÔNG THỂ SINH ĐỦ SỐ CÂU HỎI CHO CHỦ ĐỀ NÀY"
        print(notify)
    type_dict = {
        "MultipleChoice": "Tạo 1 câu hỏi trắc nghiệm MultipleChoice gồm " + str(
            number_of_answers) + " đáp án và có ít nhất 2 đáp án đúng.",
        "SingleChoice": "Tạo 1 câu hỏi trắc nghiệm singlechoice gồm " + str(
            number_of_answers) + " đáp án và có 1 đáp án đúng, " + str(number_of_answers - 1) + " đáp án sai.",
        "TrueFalse": "Tạo 1 câu hỏi trắc nghiệm TrueFalse chỉ gồm 2 loại đáp án đúng hoặc sai và có 1 đáp án đúng, 1 đáp án sai.",
    }
    mcqs = []
    for i in range(0, int(quantity)):
        if i > len(list_topic):
            continue
        if i > len(list_topic) - 1:
            continue
        genned_topic = list_topic[i]
        if genned_topic[2] == '.':
            genned_topic = genned_topic[4:]
        if genned_topic[-2] == '.':
            genned_topic = genned_topic[:-2] + genned_topic[-1]
        kq = ""
        if topic != "":
            prompt = type_dict[
                     type] + " Câu hỏi có nội dung liên quan đến " + genned_topic + " trong chủ đề \"" + topic + "\". " + \
                 bloom_dict[difficulty]
        else:
            prompt = type_dict[
                         type] + " Câu hỏi có nội dung liên quan đến " + genned_topic + ". " + \
                     bloom_dict[difficulty]
        print(f"\n\n\n----------------prompt--------------\n{prompt}")
        question = query_engine1.query(prompt)
        kq=str(question)
        eval_question = get_bloom_evaluation(kq)
        kq= str("topic: ")+str(topic)+"\n"+ str("difficulty fist: ")+str(difficulty)+"\n"+ str("eval_question_with_bloom: ")+str(eval_question)+"\n"+str(kq)
        with open("kq1.txt", "a", encoding="utf-8") as file:
            file.write(kq)
            file.write("\n\n\n")
        # mcqs.append(kq)
    # print("TẠO CÂU HỎI THÀNH CÔNG !!!")
    # print("ĐANG FORMAT CÂU HỎI ...")
    # mcqs = format_mcq(mcqs)
    # print("FORMAT CÂU HỎI THÀNH CÔNG !!!")
    # return mcqs, notify

bloom_list=["Nhớ", "Hiểu", "Áp dụng", "Phân tích", "Đánh giá", "Sáng tạo"]
topics = [
    "Khái niệm cơ sở dữ liệu là gì",
    "Sự cần thiết của các hệ cơ sở dữ liệu",
    "Mô hình kiến trúc 3 mức của cơ sở dữ liệu",
    "Mục tiêu của các hệ cơ sở dữ liệu",
    "Hệ quản trị cơ sở dữ liệu và người quản trị",
    "Tổ chức lưu trữ dữ liệu",
    "Các mô hình truy xuất dữ liệu",
    "Khái niệm về tính toàn vẹn dữ liệu",
    "Phụ thuộc hàm trong cơ sở dữ liệu",
    "Chuẩn hóa dữ liệu và các dạng chuẩn",
    "Ngôn ngữ SQL và truy vấn cơ sở dữ liệu",
    "Các phép toán trong đại số quan hệ",
    "Tối ưu hóa câu hỏi trong cơ sở dữ liệu",
    "Cơ sở dữ liệu phân tán",
    "Cơ sở dữ liệu hướng đối tượng",
    "Mô hình cơ sở dữ liệu mạng",
    "Mô hình cơ sở dữ liệu phân cấp",
    "Mô hình thực thể - liên kết",
    "Phân biệt các mô hình dữ liệu",
    "Cơ sở dữ liệu quan hệ và lý thuyết của E.F. Codd",
    "Mối quan hệ nhiều - nhiều trong cơ sở dữ liệu",
    "An toàn và bảo mật cơ sở dữ liệu",
    "Quản trị truy cập và quyền hạn trong cơ sở dữ liệu",
    "Phân loại dữ liệu trong cơ sở dữ liệu",
    "Vai trò của ánh xạ trong mô hình cơ sở dữ liệu",
    "Các chiến lược sao lưu và phục hồi dữ liệu",
    "Kiến trúc Client-Server trong cơ sở dữ liệu",
    "Ứng dụng cơ sở dữ liệu trên môi trường Internet/Intranet",
    "Cấu trúc lưu trữ vật lý của cơ sở dữ liệu",
    "Chức năng của các hệ quản trị cơ sở dữ liệu (DBMS)",
    "Hệ thống các ký hiệu biểu diễn dữ liệu",
    "Tập hợp các phép toán thao tác trên cơ sở dữ liệu",
    "Mô hình dữ liệu mạng",
    "Mô hình cơ sở dữ liệu phân cấp",
    "Mô hình cơ sở dữ liệu quan hệ",
    "Mô hình thực thể - liên kết",
    "Các đặc trưng của mô hình dữ liệu",
    "Sự ổn định trong thiết kế mô hình dữ liệu",
    "Tính đơn giản và tính dư thừa trong mô hình dữ liệu",
    "Tính đối xứng và cơ sở lý thuyết của mô hình dữ liệu",
    "Phân biệt giữa các mô hình dữ liệu",
    "Mô hình dữ liệu hướng đối tượng",
    "Mô hình dữ liệu phân tán",
    "Kiến trúc tổng quát hệ cơ sở dữ liệu 3 mức",
    "Các mô hình truy xuất dữ liệu",
    "Mô hình dữ liệu quan hệ và lý thuyết đại số quan hệ",
    "Chuẩn hóa dữ liệu và chuẩn 3NF",
    "Phương pháp khung nhìn trong tối ưu hóa câu hỏi truy vấn",
    "Quy trình tối ưu hóa câu hỏi truy vấn trong cơ sở dữ liệu",
    "Quản lý bộ đệm và bộ nhớ trong hệ quản trị cơ sở dữ liệu",
    "Quản lý các thao tác trên cơ sở dữ liệu",
    "Môi trường giao tiếp giữa người sử dụng và hệ cơ sở dữ liệu",
    "Quản lý quyền hạn truy nhập trong cơ sở dữ liệu",
    "Sự cần thiết của các hệ cơ sở dữ liệu trong quản lý thông tin",
    "Các chiến lược sao lưu và phục hồi dữ liệu",
    "Hệ thống phân tán và cơ sở dữ liệu phân tán",
    "Tính toàn vẹn của dữ liệu và các ràng buộc toàn vẹn",
    "Các phương pháp bảo vệ an toàn cơ sở dữ liệu",
    "Cấu trúc và mô hình cơ sở dữ liệu Client-Server",
    "Nguyên lý hoạt động của hệ quản trị cơ sở dữ liệu",
    "Phân loại và quản lý các quyền truy cập cơ sở dữ liệu"
]


def main():
    file_content=read_pdf_file("E:/6. Agent_MCQ_gen/MCQ-Gen/BE/CSDL giáo trình.pdf")
    import random
    for topic in topics:
        bloom = random.choice(bloom_list)
        file_path="E:/6. Agent_MCQ_gen/MCQ-Gen/BE/CSDL giáo trình.pdf"
        quantity=3
        difficulty=bloom
        number_of_answers=4
        type="SingleChoice"
        mcqGen_with_check(topic, quantity, difficulty, file_path, "", "true", type, number_of_answers)
        # print(test)